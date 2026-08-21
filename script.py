from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('BÁO CÁO DỰ ÁN KỸ THUẬT PHẦN MỀM', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.add_run('TÊN ĐỀ TÀI: ').bold = True
p.add_run('XÂY DỰNG HỆ THỐNG THI TRẮC NGHIỆM VÀ ÔN TẬP LÝ THUYẾT LÁI XE (DTS LÁI XE) DỰA TRÊN KIẾN TRÚC MICROSERVICES\n\n')
p.add_run('Người thực hiện: ').bold = True
p.add_run('[Tên của bạn]\n')
p.add_run('Mã Sinh viên/Mã NV: ').bold = True
p.add_run('[Mã số]\n')
p.add_run('Ngày hoàn thành: ').bold = True
p.add_run('Tháng 08/2026\n')

sections = [
    ("Chương 1: TỔNG QUAN DỰ ÁN", [
        ("1.1 Đặt vấn đề", "Hiện nay, nhu cầu học và thi bằng lái xe đang ngày càng tăng cao. Cục CSGT đã ban hành bộ 600 câu hỏi lý thuyết tiêu chuẩn. Việc ôn tập qua tài liệu giấy hoặc ứng dụng cũ gặp hạn chế về trải nghiệm, thiếu thống kê. Do đó, dự án 'DTS Lái xe' được xây dựng nhằm cung cấp một nền tảng trực tuyến toàn diện, giúp người dùng ôn luyện, thi thử và theo dõi năng lực một cách hiệu quả nhất."),
        ("1.2 Mục tiêu dự án", "- Về mặt nghiệp vụ: Cung cấp đầy đủ các tính năng Luyện tập theo chương, Thi thử, Chấm điểm tự động, Thống kê lịch sử tiến độ, Bảng xếp hạng. Trang quản trị để quản lý ngân hàng câu hỏi, kỳ thi và người dùng.\n- Về mặt công nghệ: Áp dụng kiến trúc Microservices hiện đại để đảm bảo tính mở rộng và chịu tải cao. Tự động hóa quy trình bằng CI/CD."),
        ("1.3 Phạm vi dự án", "- Nền tảng Web Application.\n- Phục vụ hai đối tượng chính: Người dùng cuối (User) và Quản trị viên (Admin).")
    ]),
    ("Chương 2: KHẢO SÁT VÀ PHÂN TÍCH YÊU CẦU", [
        ("2.1 Yêu cầu chức năng", "Hệ thống được chia thành hai phân hệ chính:\n- Phân hệ Học viên (User): Đăng ký, Đăng nhập, Xem danh sách khóa học, Thi thử giới hạn thời gian, Luyện tập từng chương, Theo dõi tiến độ học tập (streak), Xem bảng xếp hạng.\n- Phân hệ Quản trị viên (Admin): Quản lý danh mục câu hỏi, Cấu hình kỳ thi, Quản lý tài khoản và phân quyền, Xem thống kê toàn hệ thống."),
        ("2.2 Yêu cầu phi chức năng", "- Tính khả dụng (Availability): Hệ thống hoạt động 24/7. Nếu một service lỗi, các phần khác vẫn hoạt động (đặc thù Microservices).\n- Tính bảo mật (Security): Mật khẩu mã hóa Bcrypt. API bảo vệ bởi JWT, có cơ chế tự động Refresh Token.\n- Hiệu năng (Performance): Giao diện phản hồi mượt mà, chịu tải cao.")
    ]),
    ("Chương 3: THIẾT KẾ HỆ THỐNG VÀ KIẾN TRÚC", [
        ("3.1 Kiến trúc Tổng thể (Microservices)", "Sử dụng kiến trúc Microservices với mô hình Database-per-service (Mỗi dịch vụ sở hữu cơ sở dữ liệu độc lập).\n- API Gateway (dts-gateway): Định tuyến và kiểm soát truy cập.\n- Identity Service (dts-identity): Xác thực, cấp phát JWT, quản lý User/Role.\n- Content Builder Service (dts-content-builder): Quản lý ngân hàng câu hỏi, chương học.\n- Examination Service (dts-examination): Cấu hình kỳ thi.\n- Practice Service (dts-practice): Quản lý làm bài luyện tập.\n- Progress Service (dts-progress): Ghi nhận tiến độ học tập.\n- Result Service (dts-result): Lưu trữ kết quả thi và xuất Bảng xếp hạng."),
        ("3.2 Thiết kế Cơ sở dữ liệu", "Sử dụng Hệ quản trị cơ sở dữ liệu quan hệ PostgreSQL. Các bảng dữ liệu được phân tán để đảm bảo tính độc lập, tránh nghẽn cổ chai."),
        ("3.3 Công nghệ sử dụng", "- Backend: Java 21, Spring Boot 3.3, Spring Security.\n- Frontend: ReactJS, Next.js, Tailwind CSS, Zustand, Axios.\n- DevOps: Docker, Docker Compose, GitHub Actions.")
    ]),
    ("Chương 4: HIỆN THỰC VÀ TRIỂN KHAI CI/CD", [
        ("4.1 Quy trình Kiểm thử phần mềm (Testing)", "Viết Unit Test cho tầng Service ở toàn bộ các microservices sử dụng JUnit 5 và Mockito. Mockito giúp giả lập các Repository, đảm bảo Unit test chạy nhanh, độc lập với CSDL thật."),
        ("4.2 Tự động hóa CI/CD với GitHub Actions", "Thiết lập luồng CI/CD hoàn toàn tự động:\n- Build & Test (CI): Biên dịch code và chạy lệnh 'mvn clean test', báo cáo lỗi tự động.\n- Deploy (CD): Tự động kết nối SSH vào máy chủ VPS, tải mã nguồn và triển khai khởi động lại các Container thông qua Docker Compose."),
        ("4.3 Quản lý phiên và Bảo mật tại Frontend", "Thiết lập Axios Interceptor. Khi access token hết hạn (lỗi 401 hoặc 403), tự động gọi API Refresh Token ngầm để lấy token mới, sau đó gửi lại request ban đầu một cách hoàn toàn trong suốt với người dùng.")
    ]),
    ("Chương 5: TỔNG KẾT VÀ HƯỚNG PHÁT TRIỂN", [
        ("5.1 Kết quả đạt được", "- Hoàn thành toàn bộ các tính năng cốt lõi.\n- Hệ thống Backend và Frontend kết nối hoạt động ổn định trên mạng Internet (Deploy trên VPS).\n- Giao diện thân thiện, responsive.\n- Quy trình vận hành chuẩn hóa nhờ CI/CD."),
        ("5.2 Hướng phát triển tương lai", "- Ứng dụng Redis Cache: Đưa Bảng xếp hạng và danh sách câu hỏi vào Redis để tăng tốc độ phản hồi.\n- Giao tiếp bất đồng bộ: Áp dụng Apache Kafka hoặc RabbitMQ để các microservices trao đổi thông tin ngầm (ví dụ: tự động chấm điểm khi nộp bài thi mà không cần gọi API đồng bộ).\n- Nền tảng Mobile: Phát triển ứng dụng React Native/Flutter cho iOS và Android.")
    ])
]

for chapter_title, sections_data in sections:
    doc.add_heading(chapter_title, level=1)
    for sec_title, sec_content in sections_data:
        doc.add_heading(sec_title, level=2)
        doc.add_paragraph(sec_content)

doc.save(r'C:\Users\Dai\Desktop\Bao_Cao_DTS_Lai_Xe.docx')
print("Saved to Desktop")
